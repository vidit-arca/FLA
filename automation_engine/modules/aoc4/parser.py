import json
import os
import docx
import pdfplumber
import re

class AOC4Parser:
    def __init__(self, config_path: str):
        self.config_path = config_path
        with open(config_path, "r") as f:
            self.config = json.load(f)

    def extract_docx_text(self, path: str) -> str:
        """Extracts text from a native .docx file."""
        try:
            doc = docx.Document(path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    text += " | ".join(row_data) + "\n"
            return text
        except Exception as e:
            print(f"[!] Error reading docx {path}: {e}")
            return ""

    def extract_excel_text(self, path: str) -> str:
        """Extracts text from all sheets of a native .xlsx/.xls file."""
        import pandas as pd
        try:
            excel_file = pd.ExcelFile(path)
            text_blocks = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                # Convert the dataframe into a string format
                text_blocks.append(df.to_string(index=False, na_rep=''))
            return "\n".join(text_blocks)
        except Exception as e:
            print(f"[!] Error reading excel {path}: {e}")
            return ""

    def parse_all(self, docs: dict, ocr_outputs: dict) -> dict:
        print("[*] AOC 4 Parser initialized: Extracting full text...")
        full_text = ""
        
        # Combine all provided financial docs
        for doc_key, doc_path in docs.items():
            if not doc_path or not os.path.exists(doc_path):
                continue
                
            basename = os.path.basename(doc_path).lower()
            
            # 1. Native DOCX
            if basename.endswith(".docx") or basename.endswith(".doc"):
                print(f"[*] AOC 4 Parser: Extracting text natively from {basename}")
                full_text += self.extract_docx_text(doc_path) + "\n\n"
                
            # 1.5 Native Excel (AOC4 only)
            elif basename.endswith(".xlsx") or basename.endswith(".xls"):
                print(f"[*] AOC 4 Parser: Extracting text natively from {basename}")
                full_text += self.extract_excel_text(doc_path) + "\n\n"
                
            # 2. Native PDF (pdfplumber) or OCR Markdown
            elif basename.endswith(".pdf"):
                extracted = False
                original_basename = os.path.basename(doc_path)
                
                # Try OCR Markdown first
                if ocr_outputs and original_basename in ocr_outputs:
                    md_path = ocr_outputs[original_basename].get("md")
                    if md_path and os.path.exists(md_path):
                        print(f"[*] AOC 4 Parser: Extracting from Marker OCR Markdown for {basename}")
                        with open(md_path, "r", encoding="utf-8") as f:
                            full_text += f.read() + "\n\n"
                        extracted = True
                
                # Fallback to pdfplumber
                if not extracted:
                    print(f"[*] AOC 4 Parser: Extracting natively from PDF {basename}")
                    try:
                        with pdfplumber.open(doc_path) as pdf:
                            for page in pdf.pages:
                                text = page.extract_text()
                                if text:
                                    full_text += text + "\n"
                    except Exception as e:
                        print(f"[!] Error parsing PDF {basename}: {e}")
            
            # 3. Markdown directly
            elif basename.endswith(".md"):
                print(f"[*] AOC 4 Parser: Reading Markdown file {basename}")
                try:
                    with open(doc_path, "r", encoding="utf-8") as f:
                        full_text += f.read() + "\n\n"
                except Exception as e:
                    print(f"[!] Error reading MD {basename}: {e}")

        # Extract CIN and determine if listed
        cin_match = re.search(r'\b([LU]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6})\b', full_text, re.IGNORECASE)
        extracted_data = {
            "full_text": full_text,
            "docs": docs
        }
        
        if cin_match:
            cin_number = cin_match.group(1).upper()
            extracted_data["cin_number"] = cin_number
            if cin_number.startswith("L"):
                extracted_data["is_listed"] = "yes"
            elif cin_number.startswith("U"):
                extracted_data["is_listed"] = "no"

        return extracted_data

    def extract_financials_from_text(self, text: str, missing_keys: list) -> dict:
        """Fallback method to extract missing numeric financial data from raw OCR text/Markdown."""
        print(f"[*] AOC 4 Parser: Attempting text fallback extraction for: {missing_keys}")
        
        numeric_keywords = {
            "total_revenue": [r"total revenue", r"total income", r"revenue and other income"],
            "turnover": [r"revenue from operation", r"total turnover", r"sales turnover", r"gross turnover", r"turnover"],
            "prev_turnover": [r"previous year turnover", r"turnover.*previous year"],
            "authorised_capital": [r"authorised.*capital", r"authorized.*capital", r"authorised share capital", r"authorized share capital", r"\bauthorised\b", r"\bauthorized\b"],
            "paid_up_capital": [r"paid.?up.*capital", r"paid up share capital", r"subscribed and paid up", r"equity share capital", r"preference share capital", r"share capital"],
            "net_worth": [r"net worth", r"total equity", r"capital.*reserve.*surplus", r"reserves & surplus"],
            "prev_net_worth": [r"previous year net worth", r"net worth.*previous year"],
            "reserves_and_surplus": [r"reserves and surplus", r"reserves & surplus", r"other equity", r"retained earnings", r"reserves\s*&\s*surplus"],
            "borrowings": [r"total borrowing", r"total borrowings", r"borrowing", r"borrowings", r"long term borrowing", r"short term borrowing", r"long term borrowings", r"short term borrowings", r"secured loans", r"unsecured loans", r"loan from bank", r"loan from director", r"secured loan", r"unsecured loan"],
            "loan_to_directors_assets": [r"loans given by company to directors", r"loan given by company to directors", r"loangiven by company to directors", r"loan given by company to director", r"loan to directors"],
            "secured_loan": [r"\bsecured loan", r"\bsecured loans", r"\bsecured borrowings", r"^secured$", r"^secured\b", r"term loan taken during the year", r"^term loans?$"],
            "loan_from_directors": [r"loan from directors", r"loan from director", r"loan from shareholders", r"unsecured loan from director", r"unsecured loan taken", r"due to directors", r"loans from relatives of directors"],
            "unsecured_loan": [r"unsecured loan", r"unsecured borrowings"],
            "advance_from_customers": [r"advance from customers", r"advance from shareholders", r"security deposits"],
            # Removed dues_to_msme from text fallback because OCR consistently hallucinates numbers 
            # from the 'Trade Payables - Others' column onto the MSME line due to table merging/swapping.
            # "dues_to_msme": [r"dues to msme", r"micro and small enterprises", r"dues to micro"],
            "operating_profit": [r"operating profit", r"profit before interest", r"ebitda"],
            "net_profit_before_tax": [r"profit before tax", r"profit before exceptional items", r"pbt", r"profit.*?before.*?tax", r"profit/\s*\(loss\)\s*before\s*tax"],
            "net_profit_after_tax": [r"profit after tax", r"profit for the period", r"\bpat\b", r"profit.*?after.*?tax", r"profit/.*?\(loss\).*?for the year"],
            "rpt_monthly_remun": [r"remuneration paid to directors", r"directors remuneration", r"remuneration to directors", r"managerial remuneration", r"remuneration.*director", r"monthly remuneration", r"annual remuneration", r"appointment to any office", r"salary", r"director remuneration", r"professional charges", r"professional fees"],
            "rpt_lease": [r"lease", r"rent"],
            "rpt_sale_goods": [r"sale of goods", r"sale of material", r"sales"],
            "rpt_purchase_goods": [r"purchase of goods", r"purchase of material"],
            "loan_given_by_company": [r"loans given by company", r"loan given by company", r"loans to related parties", r"inter company loan", r"inter corporate deposit.*given", r"icd given"],
            "investments_made": [r"investments made by company", r"investment made by company", r"non.current investments", r"non current investments", r"current investments", r"investment in subsidiaries", r"investment in associates"],
            "total_loans_investments_given": [r"total loans.*given", r"loans and advances given"],
            "borrowing_defaults": [r"default in repayment", r"borrowing default"],
            "has_corporate_shareholders": [r"corporate shareholder", r"holding more than 10%", r"shareholding pattern"],
            "export_sales": [r"export of services", r"export services", r"export of service", r"exports"],
            "sitting_fees": [r"sitting fee", r"directors sitting fee", r"director sitting fee", r"sitting fees to directors"]
        }
        
        found_data = {}
        lines = text.lower().split("\n")
        
        for key in missing_keys:
            if key not in numeric_keywords:
                continue
                
            patterns = numeric_keywords[key]
            valid_number = None
            
            # Search entire document for highest priority pattern first
            for pattern in patterns:
                for idx, line in enumerate(lines):
                    if re.search(pattern, line):
                        # Skip income-tax computation lines that contain 'total income' 
                        # but are not the P&L total income (e.g., 'Gross Total Income', 'Net Total Income')
                        if key == "total_revenue" and re.search(r'gross total income|net total income|brought forward loss|deduction under chapter', line):
                            continue
                            
                        # Skip Auditor boilerplate lines that cause false positives (e.g. 'less than Rs.50 Crores')
                        if re.search(r'crores|lakhs|is exempted|notification dated|last audited financial statements|section 197', line):
                            continue
                            
                        # Skip cash flow lines for borrowings to prevent extracting 'Proceeds from borrowings'
                        if key in ["borrowings", "long_term_borrowings", "short_term_borrowings"] and re.search(r'proceeds from|repayment of|cash flow', line):
                            continue

                        # Restrict 'professional charges' to Related Party Transaction tables only
                        if key == "rpt_monthly_remun" and "professional" in pattern:
                            is_rpt_section = False
                            start_check = max(0, idx - 200)
                            end_check = min(len(lines), idx + 30)
                            for check_idx in range(start_check, end_check):
                                if re.search(r'related part(y|ies)', lines[check_idx]):
                                    is_rpt_section = True
                                    break
                            if not is_rpt_section:
                                continue

                        # Found a keyword match. Extract all numbers on this line AFTER the keyword.
                        match_obj = re.search(pattern, line)
                        search_text = line[match_obj.start():]
                        
                        # In single-line OCR layouts, a line may contain multiple labels and values (e.g., "(a) Label 1 (b) Label 2 100 200").
                        # If we matched Label 1, we must NOT pick up Label 2's values. Truncate search_text if we hit a new label marker.
                        after_kw = search_text[len(match_obj.group()):]
                        next_label_match = re.search(r'\s\([a-z]\)\s|\s\((?:i|ii|iii|iv|v|vi|vii|viii|ix|x)\)\s|\s\d+\.\s', after_kw)
                        if next_label_match:
                            search_text = search_text[:len(match_obj.group()) + next_label_match.start()]
                        
                        # Strip out quantities of shares to prevent extracting them as currency values
                        search_text = re.sub(r'\d+(?:,\d+)*\s*(?:equity\s*shares?|preference\s*shares?|shares?)', '', search_text)
                        
                        # Look ahead up to 2 lines ONLY for multi-line wrapped cells (e.g. P&L items).
                        # For balance sheet row items (where the label appears on a standalone line with
                        # no value, meaning the item is NIL/zero), do NOT look ahead - it would
                        # incorrectly pick up the next row's value!
                        has_numbers = bool(re.findall(r'(-?\s*(?:\d{1,3}(?:,\d{2,3})+|\d+)(?:\.\d+)?|\((?:\d{1,3}(?:,\d{2,3})+|\d+)(?:\.\d+)?\))', line))
                        
                        # Detect balance-sheet style rows: line contains a letter like (a), (b), (c)
                        # followed by label but no number — this means value is blank/nil in the table.
                        is_bs_row_no_value = bool(re.match(r'^\s*\(?[a-z]\)?[.)\s]', line)) and not has_numbers
                        
                        if not has_numbers and not is_bs_row_no_value and not line.strip().startswith("|"):
                            for i in range(1, 3):
                                if idx + i < len(lines):
                                    next_line = lines[idx + i]
                                    if next_line.strip().startswith("|"):
                                        break
                                    search_text += " " + next_line
                        
                        # If it's a blank balance sheet row, record 0 and move on
                        if is_bs_row_no_value:
                            valid_number = 0.0
                            print(f"    -> Found fallback '{key}' = 0.0 (Blank/Nil row: '{pattern}')")
                            break
                                    
                        # Convert standalone dashes that are table cell values (between | pipes) to '0'
                        # Only replace dashes surrounded by pipes: | - | → | 0 |
                        # This avoids replacing mathematical dashes like in 'Profit before tax (VII - VIII)'
                        search_text = re.sub(r'(?<=\|)\s*-\s*(?=\|)', ' 0 ', search_text)
                        # Also convert 'Nil' / 'NIL' to 0
                        search_text = re.sub(r'\bNil\b|\bNIL\b', ' 0 ', search_text)
                        
                        # Fix OCR artifact: Indian numbers with dots instead of commas
                        # e.g. '3.27.991' should be '3,27,991' — pattern: digit.2digits.3digits
                        search_text = re.sub(r'(\d+)\.(\d{2})\.(\d{3})', r'\1,\2,\3', search_text)
                        
                        numbers = re.findall(r'(-?\s*(?:\d{1,3}(?:,\d{2,3})+|\d+)(?:\.\d+)?|\((?:\d{1,3}(?:,\d{2,3})+|\d+)(?:\.\d+)?\))', search_text)
                        if not numbers:
                            continue
                            
                        for num_idx, num_str in enumerate(numbers):
                            clean_num = num_str.replace(",", "").replace(" ", "")
                            try:
                                if clean_num.startswith("(") and clean_num.endswith(")"):
                                    val = -float(clean_num[1:-1])
                                else:
                                    val = float(clean_num)
                            except ValueError:
                                continue
                                
                            # Filter out likely note numbers (e.g. 1, 14, 2, 6.1)
                            if val != 0 and 1 <= val < 100:
                                if "." not in clean_num and val < 50:
                                    continue
                                # If it has a decimal (e.g. 6.1) and is the very first number, check if the next number is 10x larger
                                if "." in clean_num and num_idx == 0 and len(numbers) > 1:
                                    next_num_str = numbers[1].replace(",", "").replace(" ", "")
                                    try:
                                        next_val = abs(float(next_num_str.strip("()")))
                                        if next_val > val * 10:
                                            continue
                                    except:
                                        pass
                                
                            # Filter out calendar years (1990 to 2035) for financial values (abs handles negative dashes like -2025)
                            if 1990 <= abs(val) <= 2035 and "." not in clean_num:
                                continue

                            # Skip zero if non-zero numbers appear later in the list
                            # (handles note-column '| - |' converted to 0 appearing before the actual value)
                            if val == 0:
                                remaining = numbers[num_idx + 1:]
                                has_nonzero_ahead = any(
                                    float(n.replace(",", "").replace(" ", "").strip("()")) != 0
                                    for n in remaining
                                    if n.replace(",", "").replace(" ", "").strip("()")
                                )
                                if has_nonzero_ahead:
                                    continue
                                
                            valid_number = val
                            break # Take the first valid number (usually current year)
                            
                    if valid_number is not None:
                        break # Found a valid number for this pattern, stop line search
                        
                if valid_number is not None:
                    found_data[key] = valid_number
                    print(f"    -> Found fallback '{key}' = {valid_number} (Matched: {pattern})")
                    break # We found the highest priority pattern match, skip lower priority patterns
                    
        found_data["is_subsidiary_or_holding"] = self._evaluate_holding_status(text)
        found_data["is_ind_as"] = self._evaluate_ind_as_status(text)
        return found_data

    def _evaluate_holding_status(self, full_text: str) -> str:
        """Analyzes full text to determine holding/subsidiary status."""
        text = full_text.lower()
        text = re.sub(r'\s+', ' ', text)
        
        # Step 2: High Confidence Check
        holding_phrases = [
            "holding company", "holding company:", "parent company:", "ultimate holding company",
            "immediate holding company", "subsidiaries:", "list of subsidiaries"
        ]
        sub_phrases = [
            "subsidiary company", "is a subsidiary", "the company is a subsidiary of",
            "the company is a wholly owned subsidiary"
        ]
        
        # Check holding first
        for phrase in holding_phrases:
            if phrase in text:
                print(f"[*] Holding/Subsidiary check: High confidence match on '{phrase}' (Holding)")
                return "holding"
        
        # Then check subsidiary
        for phrase in sub_phrases:
            if phrase in text:
                print(f"[*] Holding/Subsidiary check: High confidence match on '{phrase}' (Subsidiary)")
                return "subsidiary"
                
        # Step 3: Ownership Validation Check (>50%)
        ownership_matches = re.findall(r'(?:ownership|holding|subsidiary|investment).{0,50}?(\d{2,3}(?:\.\d+)?)\s*%', text)
        for match in ownership_matches:
            try:
                if float(match) > 50.0 and float(match) <= 100.0:
                    print(f"[*] Holding/Subsidiary check: Ownership > 50% match ({match}%)")
                    return "holding" # Usually means it holds an investment > 50%
            except ValueError:
                pass
                
        # Step 4: Medium Confidence Accumulation
        score = 0
        med_conf_phrases = [
            "investment in subsidiaries", "investment in associates",
            "consolidated financial statements comprise",
            "consolidated financial statements", "schedule of subsidiaries"
        ]
        for phrase in med_conf_phrases:
            if phrase in text:
                score += 20
                
        print(f"[*] Holding/Subsidiary check: Medium confidence score = {score}")
        if score >= 100:
            return "holding"
            
        return "no"

    def _evaluate_ind_as_status(self, full_text: str) -> str:
        """Analyzes full text to determine if IND AS applies based on Balance Sheet structure."""
        text = full_text.lower()
        
        # Step 1: Structural Check (Division II vs Division I format)
        # Strip out common markdown/HTML formatting that might disrupt regex
        clean_text = re.sub(r'[*_]|<b>|</b>|<br>', '', text)
        
        # Search for the first standalone row headers in the Balance Sheet table
        assets_match = re.search(r'\|\s*(?:[ivx\d\.\-\(\)]*\s*)?assets\s*\|', clean_text)
        liab_match = re.search(r'\|\s*(?:[ivx\d\.\-\(\)]*\s*)?(?:equity and liabilities|liabilities and equity|equity & liabilities)\s*\|', clean_text)
        
        if assets_match and liab_match:
            if assets_match.start() < liab_match.start():
                print(f"[*] IND AS check: Structural Match (Assets presented before Liabilities) -> Yes")
                return "yes"
            else:
                print(f"[*] IND AS check: Structural Match (Liabilities presented before Assets) -> No")
                return "no"
                
        # Step 2: Fallback Keyword Check
        text = re.sub(r'\s+', ' ', text)
        ind_as_phrases = [
            "indian accounting standard",
            "ind as",
            "companies (indian accounting standards) rules",
            "ind-as"
        ]
        
        for phrase in ind_as_phrases:
            # Add word boundaries to avoid matching "kind as" or similar
            if re.search(r'\b' + re.escape(phrase) + r'\b', text):
                print(f"[*] IND AS check: Fallback Keyword Match on '{phrase}' -> Yes")
                return "yes"
                
        return "no"

    def _detect_financials_scale(self, full_text: str) -> str:
        """Detects if the financial numbers are in Crores, Lakhs, Thousands, or Actuals."""
        if not full_text:
            return "Actuals"
        text_lower = full_text.lower()
        
        if re.search(r'(?i)(amount.*?in crores?|in crores? indian rupees|\(in crores?\))', text_lower):
            return "Crores"
        if re.search(r'(?i)(amount.*?in lakhs?|in lakhs? indian rupees|\(in lakhs?\))', text_lower):
            return "Lakhs"
        if re.search(r'(?i)(amount.*?in thousands?|in thousands? indian rupees|\(in thousands?\)|in [\'’]?000)', text_lower):
            return "Thousands"
        if re.search(r'(?i)(in hundreds?|amounts? (are )?in (?:indian )?rupees hundreds?|in (?:indian )?rupees hundreds?|amounts? (are )?in hundreds?|in hundreds? (of )?indian rupees|\(in hundreds?\))', text_lower):
            return "Hundreds"
            
        return "Actuals"
